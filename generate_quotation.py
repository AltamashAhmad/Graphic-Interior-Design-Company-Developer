from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Palette ───────────────────────────────────────────────────────────────────
DARK       = RGBColor(0x1A, 0x1A, 0x2E)
GOLD       = RGBColor(0xC9, 0xA8, 0x4C)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xF5, 0xF5, 0xF5)
MID_GREY   = RGBColor(0x88, 0x88, 0x88)
TEXT       = RGBColor(0x2C, 0x2C, 0x2C)
TH_BG      = RGBColor(0x1A, 0x1A, 0x2E)
TR_ALT     = RGBColor(0xF9, 0xF9, 0xFB)
GREEN      = RGBColor(0x1E, 0x7E, 0x4A)
RED        = RGBColor(0xC0, 0x39, 0x2B)
AMBER      = RGBColor(0xD3, 0x8B, 0x00)
BLUE_LIGHT = RGBColor(0xEB, 0xF4, 0xFF)
BLUE_TEXT  = RGBColor(0x1A, 0x50, 0x8A)
NOTE_BG    = RGBColor(0xFF, 0xF8, 0xE8)
NOTE_TEXT  = RGBColor(0x7A, 0x60, 0x10)

# ── Helpers ───────────────────────────────────────────────────────────────────

def hex_str(rgb): return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_str(rgb))
    tcPr.append(shd)

def no_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','bottom','left','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=10,
            color=TEXT, font='Calibri'):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name  = font
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    return run

def sp(para, before=0, after=0, line=None):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
    if line: para.paragraph_format.line_spacing = Pt(line)

def blank(doc, size=5):
    p = doc.add_paragraph()
    sp(p, 0, 0, size)

def gold_rule(doc):
    p = doc.add_paragraph()
    sp(p, 0, 0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'C9A84C')
    pBdr.append(bot)
    pPr.append(pBdr)

def section_heading(doc, title):
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0,0)
    set_cell_bg(c, TH_BG)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sp(p, 6, 6)
    add_run(p, f'  {title.upper()}', bold=True, size=10.5, color=GOLD)
    blank(doc, 6)

def sub_heading(doc, text):
    p = doc.add_paragraph()
    sp(p, 10, 3)
    add_run(p, text, bold=True, size=11, color=DARK)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'C9A84C')
    pBdr.append(bot)
    pPr.append(pBdr)

def build_table(doc, headers, rows, col_widths, header_bg=TH_BG):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i,col in enumerate(t.columns):
        col.width = Inches(col_widths[i])
    # header row
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_bg(c, header_bg)
        p = c.paragraphs[0]
        sp(p, 5, 5)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, h, bold=True, size=9, color=WHITE if header_bg==TH_BG else DARK)
    # data rows
    for ri,row in enumerate(rows):
        tr = t.rows[ri+1]
        bg = WHITE if ri%2==0 else TR_ALT
        for ci,cell_val in enumerate(row):
            c = tr.cells[ci]
            set_cell_bg(c, bg)
            p = c.paragraphs[0]
            sp(p, 4, 4)
            if isinstance(cell_val, list):
                for txt,bld,col in cell_val:
                    add_run(p, txt, bold=bld, size=9, color=col)
            else:
                add_run(p, str(cell_val), size=9, color=TEXT)
    blank(doc, 8)

def info_row(doc, label, value):
    p = doc.add_paragraph()
    sp(p, 2, 2)
    add_run(p, f'{label:<22}', bold=True, size=9.5, color=DARK)
    add_run(p, value, size=9.5, color=TEXT)

def note_box(doc, text, bg=NOTE_BG, fg=NOTE_TEXT):
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0,0)
    set_cell_bg(c, bg)
    p = c.paragraphs[0]
    sp(p, 7, 7)
    p.paragraph_format.left_indent  = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    add_run(p, text, size=9.5, color=fg, italic=True)
    blank(doc, 6)

def tick(doc, text, color=GREEN):
    p = doc.add_paragraph()
    sp(p, 1, 1)
    p.paragraph_format.left_indent = Inches(0.2)
    add_run(p, '✓  ', bold=True, size=10, color=color)
    add_run(p, text, size=10, color=TEXT)

def cross(doc, text):
    p = doc.add_paragraph()
    sp(p, 1, 1)
    p.paragraph_format.left_indent = Inches(0.2)
    add_run(p, '✗  ', bold=True, size=10, color=RED)
    add_run(p, text, size=10, color=TEXT)

# ══════════════════════════════════════════════════════════════════════════════
#  COVER
# ══════════════════════════════════════════════════════════════════════════════

# Dark banner
t = doc.add_table(rows=1, cols=1)
c = t.cell(0,0)
set_cell_bg(c, DARK)
no_border(c)

p = c.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(p, 20, 4)
add_run(p, 'PROJECT QUOTATION', bold=True, size=9, color=GOLD)

p2 = c.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(p2, 4, 4)
add_run(p2, 'Graphic & Interior Design App', bold=True, size=24, color=WHITE)

p3 = c.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(p3, 4, 4)
add_run(p3, 'iOS App  ·  Android App  ·  Website  ·  Admin Panel  ·  3D Configurator',
        size=10, color=GOLD)

p4 = c.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(p4, 10, 20)
add_run(p4, f'Prepared by Altamash Ahmad   ·   {date.today().strftime("%d %B %Y")}   ·   Valid for 30 days',
        size=9, color=MID_GREY)

blank(doc, 8)
gold_rule(doc)
blank(doc, 8)

# Quote info box
t2 = doc.add_table(rows=1, cols=2)
t2.alignment = WD_TABLE_ALIGNMENT.LEFT
for i,col in enumerate(t2.columns):
    col.width = Inches([3.3, 3.3][i])
lc = t2.cell(0,0)
rc = t2.cell(0,1)
no_border(lc); no_border(rc)

# Left — prepared for
for label,val in [
    ('Prepared For:', 'UAE Graphic & Interior Design Company'),
    ('Document Ref:', f'QT-2026-{date.today().strftime("%m%d")}-001'),
    ('Date:', date.today().strftime('%d %B %Y')),
    ('Valid Until:', '17 June 2026'),
]:
    p = lc.add_paragraph()
    sp(p, 2, 2)
    add_run(p, f'{label}  ', bold=True, size=9.5, color=DARK)
    add_run(p, val, size=9.5, color=TEXT)

# Right — prepared by
for label,val in [
    ('Prepared By:', 'Altamash Ahmad'),
    ('Role:', 'Full-Stack App Developer'),
    ('Email:', 'Available on request'),
    ('Phone:', 'Available on request'),
]:
    p = rc.add_paragraph()
    sp(p, 2, 2)
    add_run(p, f'{label}  ', bold=True, size=9.5, color=DARK)
    add_run(p, val, size=9.5, color=TEXT)

blank(doc, 12)

# ══════════════════════════════════════════════════════════════════════════════
#  EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

section_heading(doc, '01  —  Executive Summary')

p = doc.add_paragraph()
sp(p, 0, 6)
add_run(p, 'This document outlines the full scope, investment, timeline, and terms for the '
           'development of a multi-platform design configurator application. The platform will '
           'be delivered across three environments — iOS, Android, and Web — with a shared '
           'backend, a content management admin panel, and a 3D product customization engine.',
        size=10, color=TEXT)

p2 = doc.add_paragraph()
sp(p2, 0, 10)
add_run(p2, 'The application enables end-customers to browse interior and graphic design packages, '
            'configure 3D booth and space designs in real time, save their configurations, and '
            'submit orders directly through the app. All content — categories, designs, pricing, '
            '3D models, and banners — is fully managed through the admin panel without developer '
            'involvement after launch.',
         size=10, color=TEXT)

# Highlights box
t3 = doc.add_table(rows=1, cols=3)
t3.alignment = WD_TABLE_ALIGNMENT.LEFT
widths3 = [2.2, 2.2, 2.2]
labels3 = [
    ('55,000 AED', 'Full Project Investment\n(3 platforms + 3D)'),
    ('14–16 Weeks', 'Estimated Timeline\n(Phase 1 + Phase 2)'),
    ('3 Platforms', 'iOS · Android · Web\nAll from one codebase'),
]
for i,(big,small) in enumerate(labels3):
    c = t3.columns[i].cells[0]
    t3.columns[i].width = Inches(widths3[i])
    set_cell_bg(c, TH_BG)
    no_border(c)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p, 12, 2)
    add_run(p, big, bold=True, size=18, color=GOLD)
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p2, 2, 12)
    add_run(p2, small, size=8.5, color=MID_GREY)

blank(doc, 12)

# ══════════════════════════════════════════════════════════════════════════════
#  SCOPE OF WORK
# ══════════════════════════════════════════════════════════════════════════════

section_heading(doc, '02  —  Scope of Work')
note_box(doc, '⚡  All features listed below are fully dynamic — managed from the admin panel '
              'by your team after launch. No developer is required to add products, update '
              'pricing, change images, or manage orders.', NOTE_BG, NOTE_TEXT)

sub_heading(doc, 'A — Mobile Application (iOS & Android)')
note_box(doc, 'Built with React Native — one codebase delivers both iOS and Android apps simultaneously.',
         BLUE_LIGHT, BLUE_TEXT)

build_table(doc,
    headers=['Feature', 'Description', 'Included'],
    rows=[
        ['Splash Screen',               'Animated brand splash with logo and tagline',                              [('✓', True, GREEN)]],
        ['Home — Category Grid',        'Dynamic categories with cover images and hero banner; managed from admin', [('✓', True, GREEN)]],
        ['Design List',                 'Filterable grid of designs per category; search, filter by price/type',    [('✓', True, GREEN)]],
        ['Design Detail Page',          'Photo gallery, description, size specs, package selection (Basic/Premium/Customization)', [('✓', True, GREEN)]],
        ['3D Loading Screen',           'Pre-loads 3D model with progress indicator before configurator opens',     [('✓', True, GREEN)]],
        ['3D Configurator',             'Real-time 3D model — size swap, color palette, material swap, element toggle, orbit controls', [('✓', True, GREEN)]],
        ['Save Design Screen',          'Review configuration, select package, save to profile or proceed to checkout', [('✓', True, GREEN)]],
        ['Final Preview Screen',        'Read-only order summary with price breakdown (incl. VAT) before payment',  [('✓', True, GREEN)]],
        ['Checkout Screen',             'Contact form, order submission; WhatsApp auto-message to client team',     [('✓', True, GREEN)]],
        ['Order Confirmation Screen',   'Animated success screen, order ID, auto email confirmation to customer',   [('✓', True, GREEN)]],
        ['Profile — My Designs',        'Saved designs — tap to resume 3D configurator with all settings intact',   [('✓', True, GREEN)]],
        ['Profile — My Orders',         'Order history with live status badges (Enquiry → In Review → Confirmed)',  [('✓', True, GREEN)]],
        ['Authentication',              'Email/password sign-up + Google sign-in; guest browsing allowed',         [('✓', True, GREEN)]],
        ['Push Notifications',          'Order status change alerts (Phase 2)',                                     [('Phase 2', False, AMBER)]],
        ['Arabic Language Support',     'Full RTL layout support (if required)',                                    [('Optional', False, AMBER)]],
        ['Card Payment (PayTabs)',       'In-app card payment integration (Phase 2 or on request)',                 [('Optional', False, AMBER)]],
    ],
    col_widths=[1.9, 4.0, 0.75]
)

sub_heading(doc, 'B — Web Application')
note_box(doc, 'Built with Next.js — same features as mobile app, optimised for desktop and tablet browsers. '
              'Shares the same backend and database as the mobile app.',
         BLUE_LIGHT, BLUE_TEXT)

build_table(doc,
    headers=['Feature', 'Description', 'Included'],
    rows=[
        ['All mobile screens (web version)', 'Every screen from the mobile app, rebuilt for desktop layout', [('✓', True, GREEN)]],
        ['3D Configurator (web)',            'Three.js — same features as mobile (size/color/material/elements)', [('✓', True, GREEN)]],
        ['Side panel layout',                'Options panel fixed on right side; 3D view takes full left portion', [('✓', True, GREEN)]],
        ['Responsive design',                'Works on desktop, tablet, and mobile browser',                     [('✓', True, GREEN)]],
        ['SEO-ready pages',                  'Next.js static generation for category and design pages',          [('✓', True, GREEN)]],
    ],
    col_widths=[2.1, 3.8, 0.75]
)

sub_heading(doc, 'C — Admin Panel (Content Management)')
note_box(doc, 'A private web dashboard for the client\'s team to manage all app content — '
              'no developer needed for day-to-day operations.',
         BLUE_LIGHT, BLUE_TEXT)

build_table(doc,
    headers=['Admin Feature', 'What You Can Do'],
    rows=[
        ['Category Management',   'Add, rename, reorder, delete categories. Upload cover images and hero banners.'],
        ['Design Management',     'Add/edit/delete designs. Upload photos, PDFs, 3ds Max files. Set prices per package.'],
        ['3D Model Management',   'Upload GLB files per design per size. Assign color palettes and material textures.'],
        ['Order Management',      'View all enquiries. Update order status (In Review / Confirmed / Completed / Cancelled).'],
        ['User Management',       'View registered users, their saved designs, and order history.'],
        ['Banner Management',     'Update the home screen hero banner image and promotional banners.'],
        ['Pricing Control',       'Update package prices per design at any time. Changes reflect in app immediately.'],
    ],
    col_widths=[2.2, 4.45]
)

sub_heading(doc, 'D — Backend & Database')
build_table(doc,
    headers=['Component', 'Technology', 'Purpose'],
    rows=[
        ['Database',        'PostgreSQL (Supabase)',      'Stores all app data — users, designs, orders, categories'],
        ['Authentication',  'Supabase Auth',              'Email/password + Google sign-in + session management'],
        ['File Storage',    'Supabase Storage / S3',      'Photos, PDFs, 3D model files (GLB), texture files'],
        ['REST API',        'Supabase + Next.js API',     'All data operations between app and database'],
        ['Email Service',   'Supabase + Resend',          'Transactional emails — order confirmation, sign-up'],
        ['3D Engine (Web)', 'Three.js',                   '3D rendering, mesh control, texture swap (web)'],
        ['3D Engine (App)', 'React Three Fiber + Expo GL','3D rendering with direct GPU access (iOS/Android)'],
        ['State Management','Zustand',                    'Real-time 3D configuration state across all screens'],
    ],
    col_widths=[1.7, 2.0, 2.95]
)

# ══════════════════════════════════════════════════════════════════════════════
#  INVESTMENT
# ══════════════════════════════════════════════════════════════════════════════

section_heading(doc, '03  —  Investment (Development)')

note_box(doc, '💡  Prices below are all-inclusive development fees. They cover work across all three '
              'platforms (iOS + Android + Web) unless stated otherwise. No hidden charges.',
         NOTE_BG, NOTE_TEXT)

sub_heading(doc, 'Phase 1 — Core App (Without 3D Configurator)')

build_table(doc,
    headers=['#', 'Deliverable', 'Covers', 'Investment (AED)'],
    rows=[
        ['1', 'UI / UX Implementation',
              'All screens designed and built across iOS, Android, and Web following the approved prototype',
              '8,000'],
        ['2', 'Mobile App — iOS & Android',
              'React Native app: all screens, navigation, state management, offline-friendly',
              '14,000'],
        ['3', 'Web Application',
              'Next.js web app: all screens, SEO-ready, responsive desktop/tablet/mobile',
              '8,000'],
        ['4', 'Backend & Database',
              'Supabase setup, PostgreSQL schema, all API routes, file storage, authentication',
              '5,500'],
        ['5', 'Admin Panel',
              'Full content management dashboard: categories, designs, orders, users, banners',
              '6,500'],
        ['6', 'QA & Testing',
              'Cross-device testing on iOS (iPhone 12+), Android (Samsung A54+), Chrome, Safari',
              '3,000'],
        ['7', 'Deployment & Launch',
              'App Store submission, Google Play submission, web hosting setup, domain + SSL',
              '2,500'],
        ['', [('PHASE 1 TOTAL', True, WHITE)], '', [('47,500 AED', True, GOLD)]],
    ],
    col_widths=[0.3, 2.0, 3.45, 1.0]
)

sub_heading(doc, 'Phase 2 — 3D Customization Module')
note_box(doc, '⚡  Phase 2 is the core differentiator of this app. It can be built immediately '
              'after Phase 1 or added later as a standalone upgrade.',
         NOTE_BG, NOTE_TEXT)

build_table(doc,
    headers=['#', 'Deliverable', 'Covers', 'Investment (AED)'],
    rows=[
        ['1', '3D Engine Setup',
              'React Three Fiber + Expo GL (mobile), Three.js (web) — GLB model loading, orbit controls',
              '3,500'],
        ['2', 'Size Configurator',
              'Model swap on size selection — loads separate GLB per size variation',
              '2,500'],
        ['3', 'Color Configurator',
              'Preset palette — real-time color change on named meshes via Three.js material swap',
              '2,000'],
        ['4', 'Material Configurator',
              'Texture swap on tagged mesh groups (Wood / Metal / Glass etc.)',
              '2,500'],
        ['5', 'Element Toggle',
              'Show/hide named mesh elements (logo panel, shelving, lighting) via toggle switches',
              '1,500'],
        ['6', 'First-Person View (Look Around)',
              'Camera enters model from fixed point — user drags to look in all directions',
              '2,500'],
        ['7', 'Mobile Performance Optimisation',
              '60fps target on iPhone 12 + Samsung A54; LOD, texture compression, memory management',
              '3,500'],
        ['8', '3D Integration — Save / Preview / Checkout',
              'Configuration state passed through Screens 7 → 8 → 9 → 11 (Profile resume)',
              '1,500'],
        ['', [('PHASE 2 TOTAL', True, WHITE)], '', [('19,500 AED', True, GOLD)]],
    ],
    col_widths=[0.3, 2.1, 3.35, 1.0]
)

# Grand total table
gt = doc.add_table(rows=4, cols=2)
gt.style = 'Table Grid'
gt.alignment = WD_TABLE_ALIGNMENT.LEFT
for col in gt.columns:
    col.width = Inches(3.35)

rows_gt = [
    (LIGHT_GREY, DARK,  'Phase 1 — Core App',               '47,500 AED'),
    (LIGHT_GREY, DARK,  'Phase 2 — 3D Configurator',         '19,500 AED'),
    (LIGHT_GREY, MID_GREY, 'VAT (0% — freelance services)',  '—'),
    (TH_BG,      GOLD,  'TOTAL PROJECT INVESTMENT',          '67,000 AED'),
]
for ri,(bg,fc,label,val) in enumerate(rows_gt):
    lc = gt.rows[ri].cells[0]
    rc = gt.rows[ri].cells[1]
    set_cell_bg(lc, bg); set_cell_bg(rc, bg)
    pl = lc.paragraphs[0]; pr = rc.paragraphs[0]
    sp(pl, 6, 6); sp(pr, 6, 6)
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(pl, label, bold=(ri==3), size=10, color=fc)
    add_run(pr, val,   bold=(ri==3), size=10, color=fc)

blank(doc, 10)

# Optional add-ons
sub_heading(doc, 'Optional Add-Ons (Quoted Separately)')
build_table(doc,
    headers=['Add-On', 'Description', 'Estimate (AED)'],
    rows=[
        ['Arabic Language (RTL)',  'Full right-to-left layout; all text in Arabic + English toggle',        '6,000 – 8,000'],
        ['Card Payment (PayTabs)', 'In-app card payment gateway — Visa, Mastercard, Mada (AED)',           '4,500 – 6,000'],
        ['Push Notifications',     'Order status alerts via Expo Push Notifications',                       '2,000 – 3,000'],
        ['User Logo Upload',       'User uploads their company logo and places it on the 3D model',         '3,000 – 4,500'],
        ['Full Walkthrough (3D)',  'On-screen joystick navigation inside the 3D space (Phase 2 extension)', '5,000 – 8,000'],
        ['WhatsApp Business API',  'Automated order messages via official WhatsApp Business API (not link)', '3,000 – 4,000'],
    ],
    col_widths=[1.9, 3.7, 1.15]
)

# ══════════════════════════════════════════════════════════════════════════════
#  INFRASTRUCTURE & RUNNING COSTS
# ══════════════════════════════════════════════════════════════════════════════

section_heading(doc, '04  —  Infrastructure & Monthly Running Costs')
note_box(doc, 'These are third-party service costs paid directly to the providers after launch. '
              'They are NOT included in the development fee above. Altamash Ahmad will assist '
              'with setup and configuration of all services listed.', NOTE_BG, NOTE_TEXT)

sub_heading(doc, 'One-Time Setup Costs (paid once at launch)')
build_table(doc,
    headers=['Service', 'Provider', 'Cost', 'Notes'],
    rows=[
        ['Apple Developer Account',     'Apple',         '149 USD / year\n(~547 AED)',  'Required to publish on App Store. Renewed annually.'],
        ['Google Play Developer',       'Google',        'One-time 25 USD\n(~92 AED)',  'Required to publish on Play Store. Paid once only.'],
        ['Domain Name',                 'GoDaddy / Namecheap', '50–150 AED/year',       'e.g. yourcompany.ae or .com. Renewed annually.'],
        ['SSL Certificate',             'Let\'s Encrypt', 'Free',                       'Auto-renewing SSL — handled by hosting provider.'],
    ],
    col_widths=[2.0, 1.6, 1.5, 1.65]
)

sub_heading(doc, 'Monthly Running Costs (after launch)')
build_table(doc,
    headers=['Service', 'Purpose', 'Plan', 'Estimated Monthly Cost (AED)'],
    rows=[
        ['Supabase',
         'Database (PostgreSQL), authentication, file storage, APIs',
         'Pro — $25/month',
         '~92 AED'],
        ['Vercel',
         'Web app hosting (Next.js) — global CDN, auto-scaling',
         'Pro — $20/month',
         '~74 AED'],
        ['AWS S3 (optional)',
         'Additional file storage for large 3D GLB files and design assets',
         'Pay-as-you-go\n~$0.023/GB',
         '~50–150 AED\n(depends on file volume)'],
        ['Resend',
         'Transactional emails (order confirmations, sign-up)',
         'Free up to 3,000/month\nThen $20/month',
         '0–74 AED'],
        ['Expo EAS',
         'App build and over-the-air update service (React Native)',
         'Production — $99/month\nor per-build pricing',
         '~55–363 AED'],
        ['Apple Developer\n(annual)',
         'App Store listing (amortised monthly)',
         '149 USD/year',
         '~46 AED/month'],
        ['Domain\n(annual)',
         'Website domain renewal (amortised monthly)',
         '~100 AED/year',
         '~9 AED/month'],
    ],
    col_widths=[1.5, 2.3, 1.5, 1.45]
)

p = doc.add_paragraph()
sp(p, 0, 4)
add_run(p, 'Estimated total monthly infrastructure cost after launch:  ', size=10, color=DARK)
add_run(p, '~330–810 AED/month', bold=True, size=11, color=GOLD)

p2 = doc.add_paragraph()
sp(p2, 0, 10)
add_run(p2, '(Varies depending on user volume, storage used, and email volume. '
            'Scales up only when the business grows.)', size=9, color=MID_GREY, italic=True)

# ══════════════════════════════════════════════════════════════════════════════
#  MAINTENANCE PACKAGES
# ══════════════════════════════════════════════════════════════════════════════

section_heading(doc, '05  —  Post-Launch Maintenance Packages')
note_box(doc, 'After the app goes live, ongoing maintenance ensures the app stays secure, '
              'up-to-date with OS changes, and running smoothly. '
              'Select a package or discuss a custom arrangement.',
         NOTE_BG, NOTE_TEXT)

# Three-column comparison
t_main = doc.add_table(rows=1, cols=3)
t_main.alignment = WD_TABLE_ALIGNMENT.LEFT
pkg_data = [
    ('BASIC', '800 AED / month', DARK,   [
        ('Bug fixes and crash reports', True),
        ('Monthly security updates', True),
        ('Supabase + hosting monitoring', True),
        ('5 hours developer support/month', True),
        ('Response time: 48 hours', True),
        ('New feature development', False),
        ('Priority support', False),
        ('App Store update submissions', False),
    ]),
    ('STANDARD', '1,800 AED / month', RGBColor(0x0A,0x3D,0x62), [
        ('Everything in Basic', True),
        ('10 hours developer support/month', True),
        ('App Store + Play Store update submissions', True),
        ('Minor UI improvements', True),
        ('Response time: 24 hours', True),
        ('Monthly performance report', True),
        ('Priority support', False),
        ('Dedicated development hours (new features)', False),
    ]),
    ('PREMIUM', '3,500 AED / month', RGBColor(0x6A,0x1A,0x00), [
        ('Everything in Standard', True),
        ('20 hours developer support/month', True),
        ('Priority support — 8-hour response', True),
        ('New feature development included', True),
        ('Dedicated monthly strategy call', True),
        ('Performance optimisation included', True),
        ('3D model update assistance', True),
        ('Admin panel training sessions', True),
    ]),
]

for ci,(title,price,bg,features) in enumerate(pkg_data):
    c = t_main.columns[ci].cells[0]
    t_main.columns[ci].width = Inches(2.2)
    set_cell_bg(c, bg)
    no_border(c)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p, 10, 2)
    add_run(p, title, bold=True, size=12, color=GOLD)
    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p2, 2, 10)
    add_run(p2, price, bold=True, size=10, color=WHITE)
    for feat,included in features:
        pf = c.add_paragraph()
        sp(pf, 2, 2)
        pf.paragraph_format.left_indent = Inches(0.15)
        add_run(pf, ('✓  ' if included else '✗  '), bold=True, size=9,
                color=RGBColor(0x7E,0xDF,0xA0) if included else RGBColor(0xFF,0x7F,0x7F))
        add_run(pf, feat, size=9, color=WHITE if included else RGBColor(0xAA,0xAA,0xAA))
    p_end = c.add_paragraph()
    sp(p_end, 8, 8)

blank(doc, 10)
note_box(doc, '📋  Maintenance packages are optional and billed monthly. '
              'You may start, pause, or change packages at any time with 30 days notice. '
              'Ad-hoc support outside a package is billed at 150 AED/hour.',
         NOTE_BG, NOTE_TEXT)

# ══════════════════════════════════════════════════════════════════════════════
#  TIMELINE
# ══════════════════════════════════════════════════════════════════════════════

section_heading(doc, '06  —  Project Timeline')
note_box(doc, '⏱  Timeline begins from the date the signed agreement and first payment are received. '
              'Client-side delays (e.g. late asset delivery) will extend the timeline accordingly.',
         NOTE_BG, NOTE_TEXT)

build_table(doc,
    headers=['Phase', 'What Happens', 'Duration', 'Milestone'],
    rows=[
        ['Phase 1A\nWeeks 1–2',
         'Project kick-off, database schema design, Supabase setup, API architecture, '
         'admin panel foundation, app project scaffolding',
         '2 weeks',
         'Backend live, admin panel accessible'],
        ['Phase 1B\nWeeks 3–5',
         'Mobile app screens 1–6 (splash through 3D loading), web app home, '
         'category and design screens, authentication',
         '3 weeks',
         'Working app with real data — browsing flow complete'],
        ['Phase 1C\nWeeks 6–8',
         'Checkout flow (Screens 7–10), profile and order history (Screen 11), '
         'WhatsApp order integration, email confirmation',
         '3 weeks',
         'Full order flow working end-to-end'],
        ['Phase 1D\nWeeks 9–10',
         'QA testing across all devices, bug fixes, App Store and Play Store '
         'submission, web deployment',
         '2 weeks',
         '✅ Phase 1 Delivered — App Live'],
        ['Phase 2A\nWeeks 11–13',
         '3D engine setup, GLB model loading, orbit controls, size swap, '
         'colour palette, material swap, element toggles',
         '3 weeks',
         '3D configurator working on web and mobile'],
        ['Phase 2B\nWeeks 14–16',
         'First-person look-around mode, performance optimisation (60fps mobile), '
         'full integration with save/checkout/profile flows, QA and store update',
         '3 weeks',
         '✅ Phase 2 Delivered — 3D Feature Live'],
    ],
    col_widths=[1.2, 3.5, 1.0, 1.05]
)

# ══════════════════════════════════════════════════════════════════════════════
#  PAYMENT SCHEDULE
# ══════════════════════════════════════════════════════════════════════════════

section_heading(doc, '07  —  Payment Schedule')
note_box(doc, 'Payments are structured to align with project milestones. '
              'Work on each phase begins only after the corresponding payment is received.',
         NOTE_BG, NOTE_TEXT)

build_table(doc,
    headers=['#', 'Payment', 'When Due', 'Amount (AED)', 'Amount (if Phase 1 only)'],
    rows=[
        ['1', 'Project Kickoff Deposit',
              'Upon signing agreement',
              '20,100  (30%)',
              '14,250  (30%)'],
        ['2', 'Mid-Project Milestone',
              'Phase 1B complete — browsing flow live',
              '20,100  (30%)',
              '14,250  (30%)'],
        ['3', 'UAT / Testing Complete',
              'Phase 1D complete — app ready for store submission',
              '16,750  (25%)',
              '11,875  (25%)'],
        ['4', 'Final Delivery',
              'Phase 1 App Store live (or Phase 2 complete)',
              '10,050  (15%)',
              '7,125  (15%)'],
        ['', [('TOTAL', True, WHITE)], '',
              [('67,000 AED', True, GOLD)],
              [('47,500 AED', True, GOLD)]],
    ],
    col_widths=[0.3, 2.0, 2.0, 1.4, 1.05]
)

note_box(doc, '💳  Payment accepted via bank transfer (UAE). '
              'Invoice issued for each milestone. '
              'A 2% late payment fee applies after 7 days past due date.',
         NOTE_BG, NOTE_TEXT)

# ══════════════════════════════════════════════════════════════════════════════
#  WHAT IS INCLUDED / NOT INCLUDED
# ══════════════════════════════════════════════════════════════════════════════

section_heading(doc, '08  —  What Is & Is Not Included')

t_incl = doc.add_table(rows=1, cols=2)
t_incl.alignment = WD_TABLE_ALIGNMENT.LEFT
lc = t_incl.cell(0,0); rc = t_incl.cell(0,1)
t_incl.columns[0].width = Inches(3.3)
t_incl.columns[1].width = Inches(3.3)
no_border(lc); no_border(rc)

set_cell_bg(lc, RGBColor(0xF0,0xFB,0xF4))
set_cell_bg(rc, RGBColor(0xFD,0xF0,0xF0))

p_l = lc.paragraphs[0]; sp(p_l, 6, 4)
add_run(p_l, '✓  INCLUDED IN THIS QUOTE', bold=True, size=10, color=GREEN)

p_r = rc.paragraphs[0]; sp(p_r, 6, 4)
add_run(p_r, '✗  NOT INCLUDED', bold=True, size=10, color=RED)

included = [
    'Full source code (all platforms)',
    'iOS + Android + Web app',
    'Admin panel (content management)',
    'Backend + database + API',
    'Authentication (email + Google)',
    'WhatsApp order integration',
    'Order confirmation emails',
    'Deployment to App Store + Play Store',
    'Web deployment + SSL setup',
    '60-day post-launch bug support',
    'Admin panel usage training',
    'Git version control + code handover',
]
not_included = [
    'Apple Developer Account fee (149 USD/yr)',
    'Google Play Developer fee (25 USD one-time)',
    'Domain name purchase',
    'Supabase / Vercel / AWS monthly fees',
    'UI/UX design (app follows provided prototype)',
    '3D model files (GLB) — provided by client',
    'Photography / product images',
    'Arabic translation / copywriting',
    'PayTabs gateway account setup',
    'Ongoing content entry (designs, products)',
    'Third-party app integrations not listed here',
    'Post-warranty feature additions',
]

for txt in included:
    p = lc.add_paragraph()
    sp(p, 1, 1)
    p.paragraph_format.left_indent = Inches(0.15)
    add_run(p, '✓  ', bold=True, size=9.5, color=GREEN)
    add_run(p, txt, size=9.5, color=TEXT)

for txt in not_included:
    p = rc.add_paragraph()
    sp(p, 1, 1)
    p.paragraph_format.left_indent = Inches(0.15)
    add_run(p, '✗  ', bold=True, size=9.5, color=RED)
    add_run(p, txt, size=9.5, color=TEXT)

# Bottom padding
for cell in [lc, rc]:
    p = cell.add_paragraph(); sp(p, 6, 6)

blank(doc, 10)

# ══════════════════════════════════════════════════════════════════════════════
#  TERMS & CONDITIONS
# ══════════════════════════════════════════════════════════════════════════════

section_heading(doc, '09  —  Terms & Conditions')

tc_items = [
    ('Quotation Validity',
     'This quotation is valid for 30 days from the date of issue. Prices may change after expiry.'),
    ('Project Start',
     'Development begins upon receipt of the signed agreement and the 30% kickoff payment.'),
    ('Client Responsibilities',
     'The client must provide all required assets (logo, brand colors, images, 3D files) '
     'within agreed timeframes. Delays in asset delivery will extend the project timeline '
     'with no adjustment to the quoted price.'),
    ('Revision Policy',
     'Each project phase includes two rounds of revisions at no extra cost. Additional '
     'revision rounds are billed at 150 AED/hour.'),
    ('Scope Changes',
     'Any features not listed in this quotation constitute a scope change and will be '
     'quoted separately before work begins. No out-of-scope work is performed without '
     'written approval.'),
    ('Intellectual Property',
     'Upon receipt of the final payment, full ownership of all custom code is transferred '
     'to the client. Third-party libraries remain under their respective open-source licenses.'),
    ('Confidentiality',
     'Both parties agree to keep all project details, business information, and technical '
     'specifications confidential.'),
    ('Warranty',
     'A 60-day warranty period is included after final delivery. During this period, '
     'bugs in the delivered functionality will be fixed at no charge. This excludes new '
     'features, third-party service failures, or issues caused by client-side changes.'),
    ('Termination',
     'Either party may terminate the agreement with 14 days written notice. Work completed '
     'to date will be invoiced at the hourly rate of 150 AED/hour. Deposits are non-refundable.'),
    ('Governing Law',
     'This agreement is governed by the laws of the United Arab Emirates.'),
]

for i,(title,body) in enumerate(tc_items):
    p = doc.add_paragraph()
    sp(p, 4, 2)
    add_run(p, f'{i+1}.  {title}  — ', bold=True, size=10, color=DARK)
    add_run(p, body, size=10, color=TEXT)

blank(doc, 12)

# ══════════════════════════════════════════════════════════════════════════════
#  SIGNATURE BLOCK
# ══════════════════════════════════════════════════════════════════════════════

section_heading(doc, '10  —  Agreement & Signature')

p = doc.add_paragraph()
sp(p, 0, 8)
add_run(p, 'By signing below, both parties agree to the scope, investment, timeline, and '
           'terms outlined in this document.', size=10, color=TEXT)

t_sig = doc.add_table(rows=1, cols=2)
t_sig.alignment = WD_TABLE_ALIGNMENT.LEFT
for col in t_sig.columns:
    col.width = Inches(3.3)

for ci, (title, name) in enumerate([
    ('CLIENT',    'Authorised Signatory'),
    ('DEVELOPER', 'Altamash Ahmad'),
]):
    c = t_sig.columns[ci].cells[0]
    no_border(c)
    p = c.paragraphs[0]; sp(p, 6, 2)
    add_run(p, title, bold=True, size=9, color=GOLD)
    for label in ['Name:', 'Title:', 'Signature:', 'Date:']:
        pl = c.add_paragraph()
        sp(pl, 8, 2)
        add_run(pl, f'{label:<14}', bold=True, size=9, color=DARK)
        add_run(pl, '_' * 30, size=9, color=MID_GREY)

blank(doc, 12)
gold_rule(doc)
blank(doc, 4)

p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp(p_foot, 0, 0)
add_run(p_foot, 'Prepared by Altamash Ahmad  ·  Full-Stack App Developer  ·  UAE',
        size=8.5, color=MID_GREY)

# ── Save ──────────────────────────────────────────────────────────────────────
output = '/Users/altamashahmad/Desktop/Lance/Quotation.docx'
doc.save(output)
print(f'✅  Saved: {output}')
