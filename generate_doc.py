from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
DARK        = RGBColor(0x1A, 0x1A, 0x2E)   # deep navy
GOLD        = RGBColor(0xC9, 0xA8, 0x4C)   # gold accent
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY  = RGBColor(0xF5, 0xF5, 0xF5)
MID_GREY    = RGBColor(0x88, 0x88, 0x88)
TEXT        = RGBColor(0x2C, 0x2C, 0x2C)
NOTE_BG     = RGBColor(0xFF, 0xF8, 0xE8)
NOTE_TEXT   = RGBColor(0x7A, 0x60, 0x10)
SECTION_BG  = RGBColor(0x1A, 0x1A, 0x2E)   # same as DARK
TH_BG       = RGBColor(0x1A, 0x1A, 0x2E)
TR_ALT      = RGBColor(0xF9, 0xF9, 0xFB)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   val.get('val', 'single'))
            el.set(qn('w:sz'),    val.get('sz', '4'))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', 'auto'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=10,
            color=TEXT, font='Calibri'):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name  = font
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    return run

def para_space(para, before=0, after=0):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)

def add_blank(doc, size=4):
    p = doc.add_paragraph()
    para_space(p, 0, 0)
    p.paragraph_format.line_spacing = Pt(size)

# ── Cover header ──────────────────────────────────────────────────────────────

def add_cover(doc):
    # Dark banner
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, DARK)
    set_cell_border(cell,
        top    = {'val':'single','sz':'0','color':'1A1A2E'},
        bottom = {'val':'single','sz':'0','color':'1A1A2E'},
        left   = {'val':'single','sz':'0','color':'1A1A2E'},
        right  = {'val':'single','sz':'0','color':'1A1A2E'})
    cell._tc.get_or_add_tcPr()

    # Gold rule line in cell
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(p, 18, 4)
    add_run(p, 'APP DEVELOPMENT', bold=True, size=9,
            color=GOLD, font='Calibri')

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(p2, 2, 2)
    add_run(p2, 'Files & Information You Must Provide', bold=True,
            size=22, color=WHITE, font='Calibri')

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(p3, 6, 18)
    add_run(p3, 'Prepared by Altamash Ahmad  ·  May 13, 2026',
            size=9, color=MID_GREY, font='Calibri')

    add_blank(doc, 6)

    # Gold divider
    p4 = doc.add_paragraph()
    para_space(p4, 0, 0)
    pPr = p4._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C9A84C')
    pBdr.append(bottom)
    pPr.append(pBdr)

    add_blank(doc, 8)

# ── "What we are building" box ────────────────────────────────────────────────

def add_intro(doc):
    # Platforms
    p = doc.add_paragraph()
    para_space(p, 0, 4)
    add_run(p, 'What We Are Building', bold=True, size=13, color=DARK)

    add_blank(doc, 4)

    items = [
        ('iOS App', 'iPhone & iPad'),
        ('Android App', 'All Android devices'),
        ('Website', 'Web app — same features, same content'),
    ]
    for title, sub in items:
        p = doc.add_paragraph(style='List Bullet')
        para_space(p, 1, 1)
        add_run(p, title + '  ', bold=True, size=10, color=DARK)
        add_run(p, sub, size=10, color=TEXT)

    add_blank(doc, 6)

    # Dynamic note box
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, NOTE_BG)
    for side in ['top','bottom','left','right']:
        border_dict = {'val':'single','sz':'4','color':'C9A84C'}
        set_cell_border(cell, **{side: border_dict})

    p = cell.paragraphs[0]
    para_space(p, 8, 2)
    add_run(p, '⚡  Everything Is Dynamic — ', bold=True, size=10, color=NOTE_TEXT)
    add_run(p, 'The entire app is content-managed. Categories, designs, prices, photos, '
               '3D files, banners — all managed from the admin panel by your team. '
               'No developer needed after launch.', size=10, color=NOTE_TEXT)

    p2 = cell.add_paragraph()
    para_space(p2, 4, 8)
    add_run(p2, '📋  Working Draft — ', bold=True, size=10, color=NOTE_TEXT)
    add_run(p2, 'This document is based on discussions so far. Items may be added, '
               'removed, or updated as we continue to finalise the project scope.',
               size=10, color=NOTE_TEXT)

    add_blank(doc, 10)

# ── Section heading ───────────────────────────────────────────────────────────

def section_heading(doc, number, title):
    table = doc.add_table(rows=1, cols=1)
    cell  = table.cell(0, 0)
    set_cell_bg(cell, SECTION_BG)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_space(p, 7, 7)
    add_run(p, f'  SECTION {number}  ', bold=True, size=9,
            color=GOLD, font='Calibri')
    add_run(p, f'  {title.upper()}', bold=True, size=11,
            color=WHITE, font='Calibri')
    add_blank(doc, 6)

def sub_heading(doc, text):
    p = doc.add_paragraph()
    para_space(p, 8, 4)
    run = p.add_run(text)
    run.bold = True
    run.font.name  = 'Calibri'
    run.font.size  = Pt(11)
    run.font.color.rgb = DARK
    # underline via bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1A1A2E')
    pBdr.append(bot)
    pPr.append(pBdr)

def note_para(doc, text):
    p = doc.add_paragraph()
    para_space(p, 4, 6)
    add_run(p, '  ℹ  ', bold=True, size=9, color=GOLD)
    add_run(p, text, size=9, color=MID_GREY, italic=True)

# ── Table builder ─────────────────────────────────────────────────────────────

def add_table(doc, headers, rows_data, col_widths=None):
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows_data), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Set column widths
    if col_widths:
        for i, col in enumerate(table.columns):
            col.width = Inches(col_widths[i])

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, TH_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_space(p, 5, 5)
        add_run(p, h, bold=True, size=9, color=WHITE)

    # Data rows
    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx + 1]
        bg  = WHITE if r_idx % 2 == 0 else TR_ALT
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para_space(p, 4, 4)
            if isinstance(cell_text, list):
                for part in cell_text:
                    add_run(p, part[0], bold=part[1], size=9,
                            color=DARK if part[1] else TEXT)
            else:
                add_run(p, str(cell_text), size=9, color=TEXT)

    add_blank(doc, 8)

# ── Checklist item ────────────────────────────────────────────────────────────

def checklist_item(doc, text, sub=None):
    p = doc.add_paragraph()
    para_space(p, 1, 1)
    p.paragraph_format.left_indent = Inches(0.2)
    add_run(p, '☐  ', bold=True, size=10, color=GOLD)
    add_run(p, text, size=10, color=TEXT)
    if sub:
        add_run(p, f'  — {sub}', size=9, color=MID_GREY, italic=True)

def checklist_group(doc, title, items, note=None):
    p = doc.add_paragraph()
    para_space(p, 10, 3)
    add_run(p, title, bold=True, size=10, color=DARK)
    if note:
        add_run(p, f'  ({note})', size=9, color=MID_GREY, italic=True)
    for item, sub in items:
        checklist_item(doc, item, sub)

# ── Priority table ────────────────────────────────────────────────────────────

PRIORITY_COLORS = {
    '🔴': RGBColor(0xC0, 0x39, 0x2B),
    '🟡': RGBColor(0xD3, 0x8B, 0x00),
    '🟢': RGBColor(0x27, 0xAE, 0x60),
}

def add_priority_table(doc, items):
    table = doc.add_table(rows=1 + len(items), cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    widths = [0.6, 2.8, 3.3]
    for i, col in enumerate(table.columns):
        col.width = Inches(widths[i])

    # Header
    for i, h in enumerate(['', 'Item', 'Why It Is Urgent']):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, TH_BG)
        p = cell.paragraphs[0]
        para_space(p, 5, 5)
        add_run(p, h, bold=True, size=9, color=WHITE)

    for r_idx, (flag, item, reason) in enumerate(items):
        row = table.rows[r_idx + 1]
        bg  = WHITE if r_idx % 2 == 0 else TR_ALT

        # Flag cell
        c0 = row.cells[0]
        set_cell_bg(c0, bg)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_space(c0.paragraphs[0], 4, 4)
        add_run(c0.paragraphs[0], flag, size=12)

        # Item cell
        c1 = row.cells[1]
        set_cell_bg(c1, bg)
        para_space(c1.paragraphs[0], 4, 4)
        add_run(c1.paragraphs[0], item, bold=True, size=9, color=DARK)

        # Reason cell
        c2 = row.cells[2]
        set_cell_bg(c2, bg)
        para_space(c2.paragraphs[0], 4, 4)
        add_run(c2.paragraphs[0], reason, size=9, color=TEXT)

    add_blank(doc, 8)

# ══════════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

add_cover(doc)
add_intro(doc)

# ─── SECTION 1 — BRAND ────────────────────────────────────────────────────────
section_heading(doc, 1, 'Brand & App Identity')
note_para(doc, 'Everything in this section is needed before we can design a single screen.')

add_table(doc,
    headers=['#', 'What You Must Provide', 'File Format / Answer'],
    rows_data=[
        ['1.1', [('App name', True), (' — shown on the App Store and Google Play', False)], 'Text'],
        ['1.2', [('Logo file', True), (' — your company logo', False)], 'SVG preferred, or PNG with transparent background'],
        ['1.3', [('Brand colors', True), (' — primary and secondary colors', False)], 'Hex codes (e.g. #1A1A2E) or Pantone numbers'],
        ['1.4', [('Brand font', True), (' — font used in your brand materials', False)], 'Font file (.ttf or .otf) or font name'],
        ['1.5', [('Tagline', True), (' — short line under the logo on splash screen', False)], 'Text  (leave blank if none)'],
        ['1.6', [('Arabic support decision', True), (' — Arabic, English, or both?', False)], 'Your answer  ⚠  Affects every screen'],
    ],
    col_widths=[0.35, 3.3, 3.0]
)

# ─── SECTION 2 — CATEGORIES ───────────────────────────────────────────────────
section_heading(doc, 2, 'Category Images')
note_para(doc, 'Categories are fully dynamic — add, remove, rename, reorder anytime from the admin panel after launch. Just give us the starting set.')

add_table(doc,
    headers=['#', 'What You Must Provide', 'File Format', 'Notes'],
    rows_data=[
        ['2.1', [('Cover image', True), (' per category — shown on home screen grid', False)],
         'JPG or PNG, min 1200×800 px', 'One file per category'],
        ['2.2', [('Hero banner image', True), (' — large image above the category grid', False)],
         'JPG or PNG, min 1600×600 px', 'Can be updated from admin panel anytime'],
        ['2.3', [('Starting category list', True), (' — names exactly as shown in app', False)],
         'Text list', 'Can add / rename / reorder later from admin panel'],
        ['2.4', [('Category display order', True), (' — which appears first, second, etc.', False)],
         'Numbered list', 'Can be reordered later from admin panel'],
    ],
    col_widths=[0.35, 2.8, 2.0, 1.65]
)

sub_heading(doc, 'Starting Category List — Confirm or Edit')
note_para(doc, 'You are not locked to this list. You fully manage categories from the admin panel after launch.')
add_table(doc,
    headers=['#', 'Category Name', 'Confirmed / Edit Name Below'],
    rows_data=[
        ['1', 'Restaurants', ''],
        ['2', 'Cosmetics', ''],
        ['3', 'Booths', ''],
        ['4', 'Exhibitions', ''],
        ['5', 'Offices & Workspaces', ''],
        ['6', 'Seasonal Events', ''],
        ['+', 'Add more if needed', ''],
    ],
    col_widths=[0.35, 2.5, 3.9]
)

# ─── SECTION 3 — DESIGNS ──────────────────────────────────────────────────────
section_heading(doc, 3, 'Designs / Products')
note_para(doc, 'All design information is dynamic — managed from the admin panel. You can add, edit, remove, or reprice any design at any time after launch.')
note_para(doc, 'Two options: (A) Send us a spreadsheet and we upload for you, or (B) We give you admin access and your team enters designs directly.')

sub_heading(doc, '3A — Information (one row per design)')
add_table(doc,
    headers=['#', 'What You Must Provide', 'Format', 'Example'],
    rows_data=[
        ['3.1', [('Design name', True)],                    'Text',           '"Modern Booth A"'],
        ['3.2', [('Short description', True)],              'Text, 2–4 lines','A sleek modern exhibition booth...'],
        ['3.3', [('Category', True)],                       'From your list', '"Booths"'],
        ['3.4', [('Packages available', True)],             'Basic / Premium / Customization', 'Not every design needs all three'],
        ['3.5', [('Price per package', True), (' (AED)', False)], 'Numbers',  'Basic 500 / Premium 1,200 / Custom 2,500'],
        ['3.6', [('Available sizes', True)],                'Text list',      '3×3 m, 4×4 m, 5×8 m'],
        ['3.7', [('Available materials', True)],            'Text list',      'Wood, Metal, Glass'],
    ],
    col_widths=[0.35, 2.5, 1.9, 2.05]
)

sub_heading(doc, '3B — Files (per design)')
add_table(doc,
    headers=['#', 'What You Must Provide', 'File Format', 'Notes'],
    rows_data=[
        ['3.8', [('Product photos', True), (' — gallery images on design detail page', False)],
         'JPG or PNG, min 1500×1000 px', 'Provide 4 to 6 photos per design'],
        ['3.9', [('PDF file', True), (' — for Premium package', False)],
         'PDF', 'Technical drawings with dimensions'],
        ['3.10',[('3ds Max file', True), (' — for Premium package', False)],
         '.max file', 'Editable source file'],
    ],
    col_widths=[0.35, 2.8, 2.0, 1.65]
)

# ─── SECTION 4 — 3D FILES ─────────────────────────────────────────────────────
section_heading(doc, 4, '3D Customization Files')
note_para(doc, 'You have confirmed your team will provide all 3D files. Share this entire section with your 3D artist.')

sub_heading(doc, '4A — The 3D Model Files (GLB Format)')
note_para(doc, 'GLB is a standard 3D file format — the same model your artist creates in 3ds Max or Blender, exported in the format the app can read.')

add_table(doc,
    headers=['#', 'What Must Be Provided', 'File Format', 'Notes'],
    rows_data=[
        ['4.1', [('One 3D model file per design per size', True)],
         'GLB (.glb)', 'Each size is a SEPARATE file — see naming below'],
        ['4.2', [('Texture image files', True), (' — embedded inside the GLB', False)],
         'PNG, max 2048×2048 px', 'Must be embedded inside GLB — not separate files'],
    ],
    col_widths=[0.35, 2.8, 1.7, 1.95]
)

p = doc.add_paragraph()
para_space(p, 2, 2)
add_run(p, 'File Naming Convention — your artist must follow this exactly:', bold=True, size=9, color=DARK)

code_items = [
    'modern_booth_a_3x3.glb',
    'modern_booth_a_4x4.glb',
    'modern_booth_a_5x8.glb',
    'classic_restaurant_3x3.glb',
]
for item in code_items:
    p = doc.add_paragraph()
    para_space(p, 1, 1)
    p.paragraph_format.left_indent = Inches(0.3)
    add_run(p, item, size=9, color=RGBColor(0x1E, 0x6A, 0x4A), font='Courier New')

note_para(doc, 'If a design has 3 sizes → your artist delivers 3 separate GLB files for that design.')
add_blank(doc, 6)

sub_heading(doc, '4B — Mesh Naming Inside the GLB (Critical — For Your 3D Artist)')
note_para(doc, 'Every part users can change (color, material, show/hide) must have a unique name inside the 3D file. The artist sets these names in 3ds Max or Blender before exporting.')

add_table(doc,
    headers=['Part Type', 'Example Names to Use Inside the File'],
    rows_data=[
        ['Walls',                   'WallLeft,  WallRight,  WallBack'],
        ['Counter / Reception desk','Counter,  ReceptionDesk'],
        ['Ceiling',                 'Ceiling'],
        ['Panels / Display panels', 'Panel_Left,  Panel_Right'],
        ['Logo sign / Signage',     'LogoPanel,  Signage'],
        ['Lighting elements',       'CeilingLight,  SpotLight_1'],
        ['Flooring',                'Floor'],
        ['Any removable element',   'ExtraShelf,  DisplayStand,  LogoStand'],
    ],
    col_widths=[2.2, 4.55]
)
note_para(doc, 'Rule: If a part has no name, we cannot control it from the app. Every controllable part must be named.')
add_blank(doc, 4)

sub_heading(doc, '4C — Color Options')
note_para(doc, 'The app shows users a preset palette of colors. You define this palette.')
add_table(doc,
    headers=['#', 'What You Must Provide', 'Format', 'Notes'],
    rows_data=[
        ['4.3', [('Color palette', True), (' — list of colors users can choose from', False)],
         'Hex codes', 'Recommended 8–16 colors — your brand team decides'],
        ['4.4', [('Which parts can be colored', True), (' per design', False)],
         'Text list', 'e.g. Walls, Counter, and Ceiling can each be different colors'],
        ['4.5', [('Default color per part', True), (' — color each part starts at', False)],
         'Hex code per part', 'e.g. Walls = #FFFFFF,  Counter = #C9A84C'],
    ],
    col_widths=[0.35, 2.5, 1.3, 2.65]
)

p = doc.add_paragraph()
para_space(p, 2, 2)
add_run(p, 'Example color palette format to send us:', bold=True, size=9, color=DARK)
colors = [
    ('White',    '#FFFFFF'), ('Black',    '#1A1A1A'), ('Gold',     '#C9A84C'),
    ('Silver',   '#B0B0B0'), ('Beige',    '#F5F0E8'), ('Dark Grey','#3A3A3A'),
    ('Navy Blue','#1A2744'), ('Cream',    '#FFF8F0'),
]
for name, hex_val in colors:
    p = doc.add_paragraph()
    para_space(p, 0, 0)
    p.paragraph_format.left_indent = Inches(0.3)
    add_run(p, f'{name:<12}', size=9, color=TEXT)
    add_run(p, f'→  {hex_val}', size=9, color=RGBColor(0x1E, 0x6A, 0x4A), font='Courier New')

add_blank(doc, 8)

sub_heading(doc, '4D — Material Options')
note_para(doc, 'Materials are surface textures applied to parts of the design (e.g. wooden counter, metal walls, glass panels).')
add_table(doc,
    headers=['#', 'What You Must Provide', 'Format', 'Notes'],
    rows_data=[
        ['4.6', [('Material texture files', True), (' — one image per material', False)],
         'PNG, min 1024×1024 px, seamless', 'e.g. one file for Wood, one for Metal, one for Glass'],
        ['4.7', [('Material names', True), (' — as shown in the app', False)],
         'Text', 'e.g. "Oak Wood",  "Brushed Metal",  "Frosted Glass"'],
        ['4.8', [('Which parts can have material changed', True)],
         'Text list per design', 'e.g. Counter and Panels can switch. Walls cannot.'],
        ['4.9', [('Default material per part', True)],
         'Text', 'e.g. Counter defaults to "Oak Wood"'],
    ],
    col_widths=[0.35, 2.5, 2.0, 1.95]
)

p = doc.add_paragraph()
para_space(p, 2, 2)
add_run(p, 'Texture file naming — your artist must follow this:', bold=True, size=9, color=DARK)
for name in ['material_oak_wood.png', 'material_brushed_metal.png',
             'material_frosted_glass.png', 'material_white_marble.png']:
    p = doc.add_paragraph()
    para_space(p, 0, 0)
    p.paragraph_format.left_indent = Inches(0.3)
    add_run(p, name, size=9, color=RGBColor(0x1E, 0x6A, 0x4A), font='Courier New')

add_blank(doc, 8)

sub_heading(doc, '4E — Size Options')
add_table(doc,
    headers=['#', 'What You Must Provide', 'Format', 'Notes'],
    rows_data=[
        ['4.10', [('List of available sizes', True), (' per design', False)],
         'Text (meters)', 'e.g. 3×3 m,  4×4 m,  5×8 m'],
        ['4.11', [('One GLB file per size', True), (' per design', False)],
         'GLB (.glb)', 'Already covered in 4A — confirming for clarity'],
        ['4.12', [('Default size', True), (' — which size the design opens at', False)],
         'Text', 'e.g. "Opens in 3×3 by default"'],
    ],
    col_widths=[0.35, 2.5, 1.6, 2.35]
)

sub_heading(doc, '4F — Add / Remove Elements')
note_para(doc, 'Some parts can be toggled on or off by the user (e.g. add extra shelving, remove a counter).')
add_table(doc,
    headers=['#', 'What You Must Provide', 'Format', 'Notes'],
    rows_data=[
        ['4.13', [('List of elements that can be added or removed', True)],
         'Text list', 'e.g. Extra Shelf, Logo Stand, Ceiling Spotlights, Brochure Rack'],
        ['4.14', [('Default state per element', True), (' — ON or OFF when design opens', False)],
         'Text', 'e.g. Logo Stand = ON,  Extra Shelf = OFF'],
        ['4.15', [('These elements must be separate named meshes', True), (' in the GLB', False)],
         'Instruction for artist', 'Artist names them (e.g. ExtraShelf, LogoStand) in the file'],
    ],
    col_widths=[0.35, 2.5, 1.6, 2.35]
)

sub_heading(doc, '4G — Technical Requirements for Your 3D Artist')
note_para(doc, 'Your artist must meet these standards for the files to work in the app.')
add_table(doc,
    headers=['Requirement', 'Specification', 'Why'],
    rows_data=[
        [('Export format',    True), 'GLB only — not OBJ, FBX, STL, or GLTF+bin',          'GLB is a single self-contained file the app reads directly'],
        [('Polygon count',    True), 'Maximum 150,000 polygons per model',                   'More causes lag and overheating on mobile phones'],
        [('Texture size',     True), 'Maximum 2048×2048 pixels per texture',                 'Larger textures crash on older mobile devices'],
        [('Texture embedding',True), 'All textures embedded inside the GLB',                 'App cannot load external texture files'],
        [('Mesh naming',      True), 'Every controllable part has a unique English name',    'App controls parts by name — unnamed parts cannot be controlled'],
        [('Scale',            True), 'Real-world scale — 1 unit = 1 meter',                  'Ensures sizes display correctly'],
        [('Origin point',     True), 'Set origin to center-bottom of model',                 'Ensures model positions correctly in the app'],
        [('Test file',        True), 'Open file at: gltf-viewer.donmccurdy.com before sending', 'If it opens correctly there, it will work in the app'],
    ],
    col_widths=[1.5, 2.7, 2.6]
)

# ─── SECTION 5 — BUSINESS ─────────────────────────────────────────────────────
section_heading(doc, 5, 'Business Decisions')

add_table(doc,
    headers=['#', 'What You Must Confirm', 'Your Answer'],
    rows_data=[
        ['5.1', [('Payment method', True), (' — card inside app, or inquiry only?', False)],
         '☐  Online card payment    ☐  Inquiry only (WhatsApp / phone / email)'],
        ['5.2', [('Who manages app content', True), (' after launch?', False)],
         '☐  Technical staff    ☐  Non-technical staff'],
        ['5.3', [('Existing website', True), (' — replacing it or separate?', False)],
         'Your answer'],
        ['5.4', [('First-person walkthrough', True), (' — user steps inside the design and looks around', False)],
         '☐  Required at launch    ☐  Phase 2 is fine'],
        ['5.5', [('User logo upload', True), (' — user places their company logo on the design in 3D', False)],
         '☐  Yes, include this    ☐  No'],
    ],
    col_widths=[0.35, 2.8, 3.65]
)

# ─── DELIVERY CHECKLIST ───────────────────────────────────────────────────────
section_heading(doc, '—', 'Delivery Checklist — Complete Summary')
note_para(doc, 'Use this as your master checklist. Tick each item as files are ready to send.')

checklist_group(doc, 'Brand Files', [
    ('App name', 'text'),
    ('Logo file', 'SVG or PNG with transparent background'),
    ('Brand colors', 'hex codes'),
    ('Brand font', 'font file (.ttf / .otf) or font name'),
    ('Tagline', 'text, optional'),
    ('Arabic / English language decision', None),
])

checklist_group(doc, 'Category Files', [
    ('Starting category list with exact names', None),
    ('Category display order at launch', None),
    ('Cover image per category', 'JPG/PNG min 1200×800 px'),
    ('Hero banner image', 'JPG/PNG min 1600×600 px'),
], note='dynamic — managed from admin panel after launch')

checklist_group(doc, 'Per Design', [
    ('Design name', 'text'),
    ('Short description', '2–4 sentences'),
    ('Category', 'from your list'),
    ('Packages available', 'Basic / Premium / Customization'),
    ('Price per package', 'AED'),
    ('4 to 6 product photos', 'JPG/PNG min 1500×1000 px'),
    ('PDF file', 'Premium package only'),
    ('3ds Max file', 'Premium package only'),
], note='dynamic — managed from admin panel after launch')

checklist_group(doc, '3D Files — per design with Customization package', [
    ('GLB file per size per design', 'named correctly e.g. modern_booth_a_3x3.glb'),
    ('Color palette', 'hex codes, 8–16 colors'),
    ('Default color per part', 'hex code'),
    ('Which parts can be individually colored', 'text list'),
    ('Material texture files', 'PNG min 1024×1024 px, seamless/tileable'),
    ('Material names for the app', 'text'),
    ('Default material per part', 'text'),
    ('Which parts can have materials changed', 'text list'),
    ('List of available sizes', 'e.g. 3×3, 4×4, 5×8'),
    ('Default size', 'text'),
    ('List of add/remove elements per design', 'text'),
    ('Default state per element', 'ON or OFF'),
    ('All meshes inside GLB named correctly in English', None),
])

checklist_group(doc, 'Business Decisions', [
    ('Payment method', 'card or inquiry'),
    ('Who manages content after launch', None),
    ('First-person walkthrough', 'launch or Phase 2'),
    ('User logo upload', 'yes or no'),
])

add_blank(doc, 10)

# ─── PRIORITY TABLE ───────────────────────────────────────────────────────────
section_heading(doc, '—', 'Priority Order — Address These First')

add_priority_table(doc, [
    ('🔴', 'Logo + Brand Colors',
     'Needed before any screen can be designed'),
    ('🔴', 'Arabic support — YES or NO',
     'Affects every single screen — must decide before any UI work begins'),
    ('🔴', 'Brief your 3D artist (Section 4G)',
     '3D files are the longest lead item — the earlier the artist starts, the better'),
    ('🟡', 'Category cover images',
     'Needed for the home screen — can use placeholders temporarily'),
    ('🟡', 'Design list with photos and pricing',
     'Needed to populate the app — send spreadsheet or enter directly into admin panel'),
    ('🟢', 'Payment method decision',
     'Needed before checkout is built — not urgent for Phase 1'),
])

# ─── FOOTER ───────────────────────────────────────────────────────────────────
add_blank(doc, 8)
p = doc.add_paragraph()
para_space(p, 0, 4)
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
top = OxmlElement('w:top')
top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '6')
top.set(qn('w:space'), '1');    top.set(qn('w:color'), 'C9A84C')
pBdr.append(top)
pPr.append(pBdr)

p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Send all files and answers to: ', size=9, color=MID_GREY)
add_run(p, 'Altamash Ahmad', bold=True, size=9, color=DARK)
add_run(p, '   ·   Share files via Google Drive or WeTransfer', size=9, color=MID_GREY)

# ─── SAVE ─────────────────────────────────────────────────────────────────────
output_path = '/Users/altamashahmad/Desktop/Lance/Client-Requirements.docx'
doc.save(output_path)
print(f'✅  Saved: {output_path}')
